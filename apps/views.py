from androguard.core.bytecodes.apk import APK
from django.shortcuts import render, get_object_or_404
import imgkit
import re
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils import timezone
from docx2pdf import convert
import zipfile
from sklearn.preprocessing import StandardScaler
from torch.utils.data import DataLoader, TensorDataset
import torch
from .gcn_model import extract_features, preprocess_features, encode_features, create_feature_sequences,MYmodel
from transformers import AutoTokenizer, AutoModel
from .extract_key import analyze_urls_in_apk
from sklearn.preprocessing import LabelEncoder
from PIL import Image
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import io
from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from docx.shared import Inches
from django.views.decorators.csrf import csrf_exempt
from pyzbar.pyzbar import decode
from docx import Document
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from django.conf import settings
from lime import lime_tabular
from sklearn.model_selection import train_test_split
from django.http import FileResponse
import base64
from urllib.parse import urlparse
from androguard.misc import AnalyzeAPK
import requests
import hashlib
import xml.etree.ElementTree as ET
from OpenSSL.crypto import load_certificate, FILETYPE_ASN1, TYPE_RSA, TYPE_DSA
from cryptography.hazmat.primitives import hashes, serialization
from django.http import JsonResponse
from .models import ApkFile,WhiteBlackRecord
from datetime import datetime
import numpy as np
import joblib
import pandas as pd
from django.utils import timezone
from torch import nn
PERMISSION_DETAILS = {
    "android.permission.READ_CALENDAR": {
        "category": "危险",
        "info": "读取日历事件",
        "description": "允许应用读取用户的日历事件。"
    },
    "android.permission.WRITE_CALENDAR": {
        "category": "危险",
        "info": "写入日历事件",
        "description": "允许应用写入用户的日历事件。"
    },
    "android.permission.CAMERA": {
        "category": "危险",
        "info": "使用相机",
        "description": "允许应用访问设备的相机。"
    },
    "android.permission.READ_CONTACTS": {
        "category": "危险",
        "info": "读取联系人",
        "description": "允许应用读取用户的联系人数据。"
    },
    "android.permission.WRITE_CONTACTS": {
        "category": "危险",
        "info": "写入联系人",
        "description": "允许应用修改用户的联系人数据。"
    },
    "android.permission.GET_ACCOUNTS": {
        "category": "危险",
        "info": "获取账户信息",
        "description": "允许应用访问用户的账户列表。"
    },
    "android.permission.ACCESS_FINE_LOCATION": {
        "category": "危险",
        "info": "精确位置",
        "description": "允许应用获取精确的地理位置。"
    },
    "android.permission.ACCESS_COARSE_LOCATION": {
        "category": "危险",
        "info": "粗略位置",
        "description": "允许应用获取基于网络的粗略位置。"
    },
    "android.permission.RECORD_AUDIO": {
        "category": "危险",
        "info": "录音",
        "description": "允许应用录制音频。"
    },
    "android.permission.READ_PHONE_STATE": {
        "category": "危险",
        "info": "读取电话状态",
        "description": "允许应用访问电话状态信息。"
    },
    "android.permission.CALL_PHONE": {
        "category": "危险",
        "info": "拨打电话",
        "description": "允许应用拨打电话而无需通过拨号用户界面。"
    },
    "android.permission.READ_CALL_LOG": {
        "category": "危险",
        "info": "读取通话记录",
        "description": "允许应用读取用户的通话记录。"
    },
    "android.permission.WRITE_CALL_LOG": {
        "category": "危险",
        "info": "写入通话记录",
        "description": "允许应用修改用户的通话记录。"
    },
    "android.permission.ADD_VOICEMAIL": {
        "category": "危险",
        "info": "添加语音邮件",
        "description": "允许应用添加语音邮件。"
    },
    "android.permission.USE_SIP": {
        "category": "危险",
        "info": "使用SIP",
        "description": "允许应用使用SIP协议处理呼叫。"
    },
    "android.permission.PROCESS_OUTGOING_CALLS": {
        "category": "危险",
        "info": "处理外拨电话",
        "description": "允许应用监视、修改或中止外拨电话。"
    },
    "android.permission.BODY_SENSORS": {
        "category": "危险",
        "info": "身体传感器",
        "description": "允许应用访问身体传感器数据。"
    },
    "android.permission.SEND_SMS": {
        "category": "危险",
        "info": "发送短信",
        "description": "允许应用发送短信。"
    },
    "android.permission.RECEIVE_SMS": {
        "category": "危险",
        "info": "接收短信",
        "description": "允许应用接收和处理短信。"
    },
    "android.permission.READ_SMS": {
        "category": "危险",
        "info": "读取短信",
        "description": "允许应用读取用户的短信。"
    },
    "android.permission.RECEIVE_WAP_PUSH": {
        "category": "危险",
        "info": "接收WAP推送",
        "description": "允许应用接收WAP推送信息。"
    },
    "android.permission.RECEIVE_MMS": {
        "category": "危险",
        "info": "接收彩信",
        "description": "允许应用接收和处理彩信。"
    },
    "android.permission.READ_EXTERNAL_STORAGE": {
        "category": "危险",
        "info": "读取外部存储",
        "description": "允许应用读取外部存储中的文件。"
    },
    "android.permission.WRITE_EXTERNAL_STORAGE": {
        "category": "危险",
        "info": "写入外部存储",
        "description": "允许应用写入外部存储中的文件。"
    },
    "android.permission.INTERNET": {
        "category": "正常",
        "info": "访问网络",
        "description": "允许应用访问网络连接。"
    },
    "android.permission.ACCESS_NETWORK_STATE": {
        "category": "正常",
        "info": "访问网络状态",
        "description": "允许应用查看所有网络的状态。"
    },
    "android.permission.ACCESS_WIFI_STATE": {
        "category": "正常",
        "info": "访问Wi-Fi状态",
        "description": "允许应用查看Wi-Fi网络的状态。"
    },
    "android.permission.BLUETOOTH": {
        "category": "正常",
        "info": "使用蓝牙",
        "description": "允许应用连接到配对的蓝牙设备。"
    },
    "android.permission.BLUETOOTH_ADMIN": {
        "category": "正常",
        "info": "管理蓝牙",
        "description": "允许应用发现和配对蓝牙设备。"
    },
    "android.permission.CHANGE_NETWORK_STATE": {
        "category": "正常",
        "info": "更改网络状态",
        "description": "允许应用更改网络连接的状态。"
    },
    "android.permission.CHANGE_WIFI_STATE": {
        "category": "正常",
        "info": "更改Wi-Fi状态",
        "description": "允许应用更改Wi-Fi连接的状态。"
    },
    "android.permission.ACCESS_LOCATION_EXTRA_COMMANDS": {
        "category": "正常",
        "info": "访问额外的定位命令",
        "description": "允许应用访问额外的定位命令。"
    },
    "android.permission.WAKE_LOCK": {
        "category": "正常",
        "info": "保持唤醒状态",
        "description": "允许应用防止手机进入休眠状态。"
    },
    "android.permission.RECEIVE_BOOT_COMPLETED": {
        "category": "正常",
        "info": "接收启动完成",
        "description": "允许应用在系统完成启动后立即启动。"
    },
    "android.permission.VIBRATE": {
        "category": "正常",
        "info": "使用振动",
        "description": "允许应用控制振动器。"
    },
    "android.permission.MOUNT_UNMOUNT_FILESYSTEMS": {
        "category": "正常",
        "info": "挂载/卸载文件系统",
        "description": "允许应用挂载和卸载外部文件系统。"
    },
    "android.permission.MANAGE_DOCUMENTS": {
        "category": "正常",
        "info": "管理文档",
        "description": "允许应用管理文档存储。"
    },
    "android.permission.READ_SYNC_SETTINGS": {
        "category": "正常",
        "info": "读取同步设置",
        "description": "允许应用读取同步设置。"
    },
    "android.permission.WRITE_SYNC_SETTINGS": {
        "category": "正常",
        "info": "写入同步设置",
        "description": "允许应用写入同步设置。"
    },
    "android.permission.READ_SYNC_STATS": {
        "category": "正常",
        "info": "读取同步状态",
        "description": "允许应用读取同步状态。"
    },
    "android.permission.SET_ALARM": {
        "category": "正常",
        "info": "设置闹钟",
        "description": "允许应用设置闹钟。"
    },
    "android.permission.KILL_BACKGROUND_PROCESSES": {
        "category": "正常",
        "info": "结束后台进程",
        "description": "允许应用结束后台进程。"
    },
    "android.permission.MODIFY_AUDIO_SETTINGS": {
        "category": "正常",
        "info": "修改音频设置",
        "description": "允许应用修改全局音频设置。"
    },
    "android.permission.SYSTEM_ALERT_WINDOW": {
        "category": "系统",
        "info": "显示系统警报窗口",
        "description": "允许应用显示系统警报窗口。"
    },
    "android.permission.WRITE_SETTINGS": {
        "category": "系统",
        "info": "写入系统设置",
        "description": "允许应用修改系统设置。"
    },
    "android.permission.BIND_ACCESSIBILITY_SERVICE": {
        "category": "签名",
        "info": "绑定辅助服务",
        "description": "允许应用绑定到辅助服务。"
    },
    "android.permission.BIND_DEVICE_ADMIN": {
        "category": "签名",
        "info": "绑定设备管理员",
        "description": "允许应用绑定到设备管理器。"
    },
    "android.permission.BIND_VPN_SERVICE": {
        "category": "签名",
        "info": "绑定VPN服务",
        "description": "允许应用绑定到VPN服务。"
    }
    # 请根据需要补充更多权限
}
API_CATEGORIES = {
    "Landroid/app/": "Application",
    "Landroid/content/": "Content",
    "Landroid/view/": "View",
    "Landroid/widget/": "Widget",
    "Landroid/os/": "OS",
    "Landroid/provider/": "Provider",
    "Ljava/io/": "IO",
    "Ljava/lang/": "Lang",
    "Ljava/util/": "Util",
    "Ljava/net/": "Net",
}
def home(request):
    return render(request, 'home.html')
def about_us(request):
    return render(request, 'about_us.html')
def black_white_filter(request):
    if request.method == 'POST':
        if 'apk_file' in request.FILES:
            return handle_apk_file_filter(request)
        elif 'qr_code' in request.FILES:
            return handle_qr_code_filter(request)
        elif 'download_link' in request.POST:
            return handle_download_link_filter(request)
    return render(request, 'black_white_filter.html')
def handle_apk_file_filter(request):
    apk_file = request.FILES['apk_file']
    file_name_without_extension = os.path.splitext(apk_file.name)[0]
    apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', file_name_without_extension)

    if not os.path.exists(apk_dir_path):
        os.makedirs(apk_dir_path, exist_ok=True)
    file_path = os.path.join(apk_dir_path, apk_file.name)
    if os.path.exists(file_path):
        os.remove(file_path)
    fs = FileSystemStorage(location=apk_dir_path)
    filename = fs.save(apk_file.name, apk_file)
    file_path = os.path.join(fs.location, filename)

    return process_apk_file_filter(file_path, apk_file.name, request)


def handle_qr_code_filter(request):
    if 'qr_code' not in request.FILES:
        return JsonResponse({'status': 'error', 'message': '没有上传文件。'})

    qr_code_file = request.FILES['qr_code']
    try:
        image = Image.open(qr_code_file)
        decoded_objects = decode(image)
        if not decoded_objects:
            return JsonResponse({'status': 'error', 'message': '无法解码二维码，请确保上传的是有效的二维码图片。'})

        download_link = decoded_objects[0].data.decode('utf-8')
        return download_apk_from_link_filter(download_link, request)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'处理二维码时出错: {str(e)}'})


def handle_download_link_filter(request):
    download_link = request.POST['download_link']
    return download_apk_from_link_filter(download_link, request)

def download_apk_from_link_filter(download_link, request):
    try:
        response = requests.get(download_link, proxies={'http': None, 'https': None}, timeout=10)  # 设置一个明确的超时时间
        response.raise_for_status()

        content_type = response.headers.get('content-type', '')
        if 'application/vnd.android.package-archive' not in content_type:
            return JsonResponse({'status': 'error', 'message': '下载链接不指向有效的APK文件。'})

        url_path = urlparse(download_link).path
        file_name = os.path.basename(url_path)
        if not file_name.endswith('.apk'):
            return JsonResponse({'status': 'error', 'message': '下载文件不是一个APK。'})

        apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0])
        os.makedirs(apk_dir_path, exist_ok=True)
        file_path = os.path.join(apk_dir_path, file_name)

        with open(file_path, 'wb') as apk_file:
            apk_file.write(response.content)

        return process_apk_file_filter(file_path, file_name, request)
    except Timeout:
        return JsonResponse({'status': 'error', 'message': '连接超时，请检查网络连接或链接是否有效。'})
    except ConnectionError:
        return JsonResponse({'status': 'error', 'message': '无法连接到下载服务器，请稍后再试。'})
    except requests.exceptions.RequestException as e:
        return JsonResponse({'status': 'error', 'message': '下载失败，请检查链接或稍后重试。'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': '处理过程中出现未知错误。'})
def process_apk_file_filter(file_path, file_name, request):
    try:
        sha256_hash = calculate_hash(file_path, 'sha256')
        record = WhiteBlackRecord.objects.filter(sha256_value=sha256_hash).first()
        if record:
            return JsonResponse({'status': 'success', 'message': f'文件类别: {record.category}'})
        else:
            # 如果未找到记录，返回过滤未完成信息
            return JsonResponse({'status': 'error', 'message': '过滤未完成，请移步其他操作核验。'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})
def static_ana(request, file_name):
    apk_record = get_object_or_404(ApkFile, file_name=file_name)
    analysis_result = apk_record.static_analysis_result
    return render(request, 'static_ana.html', {'analysis_result': analysis_result})

def recent_scan(request):
    recent_apks = ApkFile.objects.order_by('-upload_time')
    return render(request, 'recent_scan.html', {'recent_apks': recent_apks})


@csrf_exempt
def upload_model_predict(request):
    if request.method == 'POST':
        if 'apk_file' in request.FILES:
            return handle_apk_file_upload(request)
        elif 'qr_code' in request.FILES:
            return handle_qr_code_upload(request)
        elif 'download_link' in request.POST:
            return handle_download_link_upload(request)
    return render(request, 'upload_model_predict.html')


def handle_apk_file_upload(request):
    apk_file = request.FILES['apk_file']
    file_name_without_extension = os.path.splitext(apk_file.name)[0]
    apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', file_name_without_extension)

    if not os.path.exists(apk_dir_path):
        os.makedirs(apk_dir_path, exist_ok=True)
    file_path = os.path.join(apk_dir_path, apk_file.name)
    if os.path.exists(file_path):
        os.remove(file_path)
    fs = FileSystemStorage(location=apk_dir_path)
    filename = fs.save(apk_file.name, apk_file)
    file_path = os.path.join(fs.location, filename)

    return process_apk_file(file_path, apk_file.name, request)


def handle_qr_code_upload(request):
    if 'qr_code' not in request.FILES:
        return JsonResponse({'status': 'error', 'message': '没有上传文件。'})

    qr_code_file = request.FILES['qr_code']
    try:
        image = Image.open(qr_code_file)
        decoded_objects = decode(image)
        if not decoded_objects:
            return JsonResponse({'status': 'error', 'message': '无法解码二维码，请确保上传的是有效的二维码图片。'})

        download_link = decoded_objects[0].data.decode('utf-8')
        return download_apk_from_link(download_link, request)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'处理二维码时出错: {str(e)}'})


def handle_download_link_upload(request):
    download_link = request.POST.get('download_link')
    if not download_link:
        return JsonResponse({'status': 'error', 'message': '下载链接不能为空。'})
    return download_apk_from_link(download_link, request)
def download_apk_from_link(download_link, request):
    try:
        response = requests.get(download_link, proxies={'http': None, 'https': None}, timeout=10)  # 设置一个明确的超时时间
        response.raise_for_status()

        content_type = response.headers.get('content-type', '')
        if 'application/vnd.android.package-archive' not in content_type:
            return JsonResponse({'status': 'error', 'message': '下载链接不指向有效的APK文件。'})

        url_path = urlparse(download_link).path
        file_name = os.path.basename(url_path)
        if not file_name.endswith('.apk'):
            return JsonResponse({'status': 'error', 'message': '下载文件不是一个APK。'})

        apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0])
        os.makedirs(apk_dir_path, exist_ok=True)
        file_path = os.path.join(apk_dir_path, file_name)

        with open(file_path, 'wb') as apk_file:
            apk_file.write(response.content)

        return process_apk_file(file_path, file_name, request)
    except Timeout:
        return JsonResponse({'status': 'error', 'message': '连接超时，请检查网络连接或链接是否有效。'})
    except ConnectionError:
        return JsonResponse({'status': 'error', 'message': '无法连接到下载服务器，请稍后再试。'})
    except requests.exceptions.RequestException as e:
        return JsonResponse({'status': 'error', 'message': '下载失败，请检查链接或稍后重试。'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': '处理过程中出现未知错误。'})
def process_apk_file(file_path, file_name, request):
    try:
        md5_hash = calculate_hash(file_path, 'md5')

        apk_record, created = ApkFile.objects.get_or_create(md5=md5_hash, defaults={
            'file_name': file_name,
            'file_path': file_path,
            'upload_time': timezone.now(),
        })

        if not created:
            apk_record.file_name = file_name
            apk_record.file_path = file_path
            apk_record.upload_time = timezone.now()
            apk_record.save()

        file_path_csv = os.path.join(settings.BASE_DIR, 'secureapp', 'data', 'apk_features_train.csv')
        file_path_multi_csv = os.path.join(settings.BASE_DIR, 'secureapp', 'data', 'features.csv')

        # Load data
        data_multi = pd.read_csv(file_path_multi_csv, encoding='latin1')
        df = pd.read_csv(file_path_csv, encoding='latin1')

        # Prepare multi data
        feature_columns_multi = [col for col in data_multi.columns if col not in ['label', 'apk_file']]
        all_permissions_multi = feature_columns_multi
        df_multi = data_multi.drop('apk_file', axis=1)
        X_multi = df_multi.drop('label', axis=1)
        y_multi = df_multi['label']

        label_encoder = LabelEncoder()
        y_encoded = label_encoder.fit_transform(y_multi)

        X_train_multi, X_test_multi, y_train_multi, y_test_multi = train_test_split(X_multi, y_encoded, test_size=0.2, random_state=42)

        # Prepare main data
        df = df.drop('apk_file', axis=1)
        X = df.drop('label', axis=1)
        y = df['label']
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

        explainer = lime_tabular.LimeTabularExplainer(
            training_data=np.array(X_train),
            feature_names=X_train.columns,
            class_names=['white', 'scam'],
            mode='classification'
        )

        model_file_path = os.path.join(settings.BASE_DIR, 'secureapp', 'models', 'model.pkl')
        model = joblib.load(model_file_path)
        rf_clf_file_path = os.path.join(settings.BASE_DIR, 'secureapp', 'models', 'random_forest_model.pkl')
        rf_clf = joblib.load(rf_clf_file_path)
        label_encoder_file_path = os.path.join(settings.BASE_DIR, 'secureapp', 'models', 'label_encoder.pkl')
        label_encoder = joblib.load(label_encoder_file_path)
        feature_columns = [col for col in df.columns if col not in ['label', 'apk_file']]
        all_permissions = feature_columns

        new_permissions = extract_permissions(file_path)

        # Ensure new feature vectors match the model's expected feature dimensions
        new_feature_vector = [1 if permission in new_permissions else 0 for permission in all_permissions]
        X_new = np.array(new_feature_vector).reshape(1, -1)
        if X_new.shape[1] != len(all_permissions):
            return JsonResponse({'status': 'error', 'message': f'Feature vector length mismatch: expected {len(all_permissions)}, got {X_new.shape[1]}'})

        prediction = model.predict(X_new)
        prediction = model.predict(X_new)
        prediction_proba = model.predict_proba(X_new)
        confidence = prediction_proba.max()

        new_feature_vector_multi = [1 if permission in new_permissions else 0 for permission in all_permissions_multi]
        X_new_multi = np.array(new_feature_vector_multi).reshape(1, -1)
        if X_new_multi.shape[1] != len(all_permissions_multi):
            return JsonResponse({'status': 'error', 'message': f'Feature vector length mismatch: expected {len(all_permissions_multi)}, got {X_new_multi.shape[1]}'})

        exp = explainer.explain_instance(
            data_row=np.array(new_feature_vector),
            predict_fn=model.predict_proba
        )
        lime_output_relative_path = os.path.join('apk_files', os.path.splitext(file_name)[0], 'lime_explanation.html')
        lime_output_path = os.path.join(settings.MEDIA_ROOT, lime_output_relative_path)
        exp.save_to_file(lime_output_path)

        predicted_label_multi = rf_clf.predict(X_new_multi)
        predicted_class_multi = label_encoder.inverse_transform(predicted_label_multi)
        explainer_multi = lime_tabular.LimeTabularExplainer(
            X_train_multi.values,
            feature_names=X_multi.columns,
            class_names=label_encoder.classes_,
            discretize_continuous=True
        )

        explanation = explainer_multi.explain_instance(X_new_multi[0], rf_clf.predict_proba, num_features=10, labels=range(4))
        lime_output_relative_path_multi = os.path.join('apk_files', os.path.splitext(file_name)[0], 'lime_explanation_multi.html')
        lime_output_path_multi = os.path.join(settings.MEDIA_ROOT, lime_output_relative_path_multi)
        explanation.save_to_file(lime_output_path_multi)

        apk_record.prediction_result = int(prediction[0])
        apk_record.confidence=confidence
        apk_record.lime_explanation_path = lime_output_relative_path  # 保存相对路径
        apk_record.lime_explanation_path_multi = lime_output_relative_path_multi
        features = extract_features(file_path)
        preprocessed_features = preprocess_features(features)
        local_model_path = './bert-base-uncased'
        tokenizer = AutoTokenizer.from_pretrained(local_model_path)
        bert_model = AutoModel.from_pretrained(local_model_path)
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        encoded_features = encode_features(bert_model, tokenizer, preprocessed_features, device)
        feature_sequences = create_feature_sequences(encoded_features)
        df_new = pd.DataFrame(feature_sequences)
        features_new = df_new.values
        scaler = StandardScaler()
        scaled_features_new = scaler.fit_transform(features_new)
        features_tensor_new = torch.tensor(scaled_features_new, dtype=torch.float32,device=device)
        features_tensor = features_tensor_new.view(-1, 768, 1)
        test_dataset_new = TensorDataset(features_tensor)
        test_loader_new = DataLoader(test_dataset_new, batch_size=1, shuffle=False)
        gcn_model_file_path = os.path.join(settings.BASE_DIR, 'secureapp', 'models', '0.pth')
        gcn_model=torch.load(gcn_model_file_path)
        gcn_model.eval()
        gcn_model.to(device)
        gcn_prediction = None
        gcn_confidence = None
        with torch.no_grad():
            for inputs in test_loader_new:
                inputs = inputs[0]
                outputs = gcn_model(inputs)
                probabilities = torch.nn.functional.softmax(outputs, dim=1)
                _, predicted = torch.max(outputs, 1)
                gcn_prediction = predicted.item()
                gcn_confidence = probabilities.max().item()
        apk_record.gcn_prediction_result = gcn_prediction
        apk_record.gcn_confidence = gcn_confidence
        apk_record.save()
        formatted_time = apk_record.upload_time.strftime('%Y-%m-%d %H:%M:%S')
        prediction_result = {
            'file_name': file_name,
            'prediction': int(prediction[0]),
            'confidence': confidence,
            'gcn_prediction': gcn_prediction,
            'gcn_confidence': gcn_confidence,
            'lime_output_path': settings.MEDIA_URL + lime_output_relative_path,
            'lime_output_path_multi': settings.MEDIA_URL + lime_output_relative_path_multi,
            'md5': md5_hash,
            'scan_time': formatted_time
        }
        request.session['prediction_result'] = prediction_result
        word_file_path,pdf_file_path = generate_file_model(prediction_result, file_name)
        apk_record.word_report_path_model = word_file_path
        apk_record.pdf_report_path_model = pdf_file_path
        apk_record.save()
        result_url = f"/model_predict/{file_name}/"
        return JsonResponse({'status': 'success', 'redirect_url': result_url})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})


def download_model(request):
    zip_file_path = os.path.join(settings.MEDIA_ROOT, 'model.zip')
    response = FileResponse(open(zip_file_path, 'rb'))
    response['Content-Disposition'] = f'attachment; filename="{os.path.basename(zip_file_path)}"'
    return response
def download_project(request):
    zip_file_path = os.path.join(settings.MEDIA_ROOT, 'SecurityGuard.zip')
    response = FileResponse(open(zip_file_path, 'rb'))
    response['Content-Disposition'] = f'attachment; filename="{os.path.basename(zip_file_path)}"'
    return response
from django.http import HttpResponse, Http404
import pytz
def model_predict(request, apk_file_name):
    full_file_name = apk_file_name
    apk_record = get_object_or_404(ApkFile, file_name=full_file_name)
    lime_output_path = apk_record.lime_explanation_path
    lime_output_path_multi = apk_record.lime_explanation_path_multi
    beijing_tz = pytz.timezone('Asia/Shanghai')
    beijing_time = apk_record.upload_time.astimezone(beijing_tz)
    formatted_time = beijing_time.strftime('%Y-%m-%d %H:%M:%S')
    prediction_result = {
        'file_name': apk_record.file_name,
        'md5': apk_record.md5,
        'scan_time':formatted_time,
        'prediction': apk_record.prediction_result,
        'confidence':apk_record.confidence,
        'gcn_prediction': apk_record.gcn_prediction_result,
        'gcn_confidence': apk_record.gcn_confidence,
        'lime_output_path': os.path.join(settings.MEDIA_URL, lime_output_path),
        'lime_output_path_multi': os.path.join(settings.MEDIA_URL, lime_output_path_multi)# 转换为相对路径
    }
    return render(request, 'model_predict.html', {'prediction_result': prediction_result})


@csrf_exempt
def delete_apk(request, apk_file_name):
    if request.method == 'DELETE':
        try:
            apk_record = get_object_or_404(ApkFile, file_name=apk_file_name)
            apk_record.delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'})
def extract_permissions(apk_path):
    try:
        apk = APK(apk_path)
        return apk.get_permissions()
    except Exception as e:
        print(f"Error processing {apk_path}: {e}")
        return []
def dynamic_ana(request):
    return render(request, 'dynamic_ana.html')
def upload_and_dynamic_ana(request):
    return render(request, 'upload_and_dynamic_ana.html')
@csrf_exempt
def upload_and_static_ana(request):
    if request.method == 'POST':
        if 'apk_file' in request.FILES:
            return handle_apk_file_upload_static(request)
        elif 'qr_code' in request.FILES:
            return handle_qr_code_upload_static(request)
        elif 'download_link' in request.POST:
            return handle_download_link_upload_static(request)
    return render(request, 'upload_and_static_ana.html')
@csrf_exempt
def rescan_apk(request):
    if request.method == 'POST':
        try:
            apk_file_name = request.POST['apk_file_name']
            apk_record = get_object_or_404(ApkFile, file_name=apk_file_name)
            file_path = apk_record.file_path
            analysis_result = analyze_apk(file_path)
            apk_record.package_name = analysis_result.get('package')
            apk_record.app_name = analysis_result.get('app_name')
            apk_record.version_name = analysis_result.get('version_name')
            apk_record.static_analyzed = True
            apk_record.save()
            # 更新session中的分析结果
            request.session['analysis_result'] = analysis_result
            response_data = {
                'status': 'success',
                'redirect_url': request.build_absolute_uri(f'/static_ana/{apk_file_name}')
            }
            return JsonResponse(response_data)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'})

@csrf_exempt
def static_predict_apk(request):
    if request.method == 'POST':
        try:
            apk_file_name = request.POST['apk_file_name']
            apk_record = get_object_or_404(ApkFile, file_name=apk_file_name)
            file_path = apk_record.file_path

            # 执行预测逻辑
            file_path_csv = os.path.join(settings.BASE_DIR, 'secureapp', 'data', 'apk_features_train.csv')
            df = pd.read_csv(file_path_csv, encoding='latin1')
            df = df.drop('apk_file', axis=1)
            X = df.drop('label', axis=1)
            y = df['label']
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
            explainer = lime_tabular.LimeTabularExplainer(
                training_data=np.array(X_train),
                feature_names=X_train.columns,
                class_names=['white', 'scam'],
                mode='classification'
            )

            model_file_path = os.path.join(settings.BASE_DIR, 'secureapp', 'models', 'model.pkl')
            model = joblib.load(model_file_path)

            feature_columns = [col for col in df.columns if col not in ['label', 'apk_file']]
            all_permissions = feature_columns
            new_permissions = extract_permissions(file_path)
            new_feature_vector = [1 if permission in new_permissions else 0 for permission in all_permissions]
            X_new = np.array(new_feature_vector).reshape(1, -1)
            prediction = model.predict(X_new)
            prediction_proba = model.predict_proba(X_new)
            confidence = prediction_proba.max()

            exp = explainer.explain_instance(
                data_row=np.array(new_feature_vector),
                predict_fn=model.predict_proba
            )
            lime_output_relative_path = os.path.join('apk_files', os.path.splitext(apk_file_name)[0], 'lime_explanation.html')
            lime_output_path = os.path.join(settings.MEDIA_ROOT, lime_output_relative_path)
            exp.save_to_file(lime_output_path)

            apk_record.prediction_result = int(prediction[0])
            apk_record.lime_explanation_path = lime_output_relative_path
            apk_record.save()

            prediction_result = {
                'file_name': apk_file_name,
                'prediction': int(prediction[0]),
                'lime_output_path': settings.MEDIA_URL + lime_output_relative_path
            }
            request.session['prediction_result'] = prediction_result

            response_data = {
                'status': 'success',
                'redirect_url': request.build_absolute_uri('/model_predict/' + apk_file_name)
            }
            return JsonResponse(response_data)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'})
def handle_apk_file_upload_static(request):
    apk_file = request.FILES['apk_file']
    file_name_without_extension = os.path.splitext(apk_file.name)[0]
    apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', file_name_without_extension)

    if not os.path.exists(apk_dir_path):
        os.makedirs(apk_dir_path, exist_ok=True)
    file_path = os.path.join(apk_dir_path, apk_file.name)
    if os.path.exists(file_path):
        os.remove(file_path)
    fs = FileSystemStorage(location=apk_dir_path)
    filename = fs.save(apk_file.name, apk_file)
    file_path = os.path.join(fs.location, filename)

    return process_apk_file_static(file_path, apk_file.name, request)
def handle_qr_code_upload_static(request):
    if 'qr_code' not in request.FILES:
        return JsonResponse({'status': 'error', 'message': '没有上传文件。'})

    qr_code_file = request.FILES['qr_code']
    try:
        image = Image.open(qr_code_file)
        decoded_objects = decode(image)
        if not decoded_objects:
            return JsonResponse({'status': 'error', 'message': '无法解码二维码，请确保上传的是有效的二维码图片。'})

        download_link = decoded_objects[0].data.decode('utf-8')
        return download_apk_from_link_static(download_link, request)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'处理二维码时出错: {str(e)}'})

def handle_download_link_upload_static(request):
    download_link = request.POST['download_link']
    return download_apk_from_link_static(download_link, request)

from requests.exceptions import ConnectionError, Timeout
def download_apk_from_link_static(download_link, request):
    try:
        response = requests.get(download_link, proxies={'http': None, 'https': None}, timeout=10)  # 设置一个明确的超时时间
        response.raise_for_status()

        content_type = response.headers.get('content-type', '')
        if 'application/vnd.android.package-archive' not in content_type:
            return JsonResponse({'status': 'error', 'message': '下载链接不指向有效的APK文件。'})

        url_path = urlparse(download_link).path
        file_name = os.path.basename(url_path)
        if not file_name.endswith('.apk'):
            return JsonResponse({'status': 'error', 'message': '下载文件不是一个APK。'})

        apk_dir_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0])
        os.makedirs(apk_dir_path, exist_ok=True)
        file_path = os.path.join(apk_dir_path, file_name)

        with open(file_path, 'wb') as apk_file:
            apk_file.write(response.content)

        return process_apk_file_static(file_path, file_name, request)
    except Timeout:
        return JsonResponse({'status': 'error', 'message': '连接超时，请检查网络连接或链接是否有效。'})
    except ConnectionError:
        return JsonResponse({'status': 'error', 'message': '无法连接到下载服务器，请稍后再试。'})
    except requests.exceptions.RequestException as e:
        return JsonResponse({'status': 'error', 'message': '下载失败，请检查链接或稍后重试。'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': '处理过程中出现未知错误。'})

def process_apk_file_static(file_path, file_name, request):
    try:
        md5_hash = calculate_hash(file_path, 'md5')

        apk_record, created = ApkFile.objects.get_or_create(md5=md5_hash, defaults={
            'file_name': file_name,
            'file_path': file_path,
            'upload_time': datetime.now(),
        })

        if not created:
            apk_record.file_name = file_name
            apk_record.file_path = file_path
            apk_record.upload_time = datetime.now()
            apk_record.save()

        analysis_result = analyze_apk(file_path)
        word_file_path,pdf_file_path = generate_file_static(analysis_result, file_name)
        apk_record.word_report_path_static = word_file_path
        apk_record.pdf_report_path_static = pdf_file_path
        apk_record.package_name = analysis_result.get('package')
        icon_data = analysis_result.get('app_icon_base64')
        apk_record.app_name = analysis_result.get('app_name')
        apk_record.version_name = analysis_result.get('version_name')
        apk_record.static_analysis_result = analysis_result
        if icon_data:
            apk_record.app_icon_base64 = icon_data
        apk_record.static_analyzed = True
        apk_record.save()
        # 存储分析结果到 session
        request.session['analysis_result'] = analysis_result
        # 返回包含重定向URL的JSON响应
        response_data = {
            'status': 'success',
            'redirect_url': request.build_absolute_uri(f'/static_ana/{file_name}/')
        }
        return JsonResponse(response_data)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)})

    return render(request, 'upload_and_static_ana.html')

def analyze_apk(file_path):
    analysis_result = {}
    try:
        apk = APK(file_path)
        analysis_result['file_name'] = os.path.basename(file_path)
        analysis_result['size'] = format_size_to_mb(os.path.getsize(file_path))
        analysis_result['md5'] = calculate_hash(file_path, 'md5')
        analysis_result['sha1'] = calculate_hash(file_path, 'sha1')
        analysis_result['sha256'] = calculate_hash(file_path, 'sha256')

        analysis_result['package'] = apk.get_package() or None
        analysis_result['main_activity'] = apk.get_main_activity() or None
        analysis_result['activities'] = apk.get_activities() or []
        analysis_result['services'] = apk.get_services() or []
        analysis_result['receivers'] = apk.get_receivers() or []
        analysis_result['providers'] = apk.get_providers() or []
        analysis_result['permissions'] = []
        for perm in (apk.get_permissions() or []):
            details = PERMISSION_DETAILS.get(perm, {"category": "正常", "info": "无详细信息", "description": "无描述"})
            analysis_result['permissions'].append({'Permission': perm,'Status': details["category"],'Info': details["info"],'Description': details["description"]})
        analysis_result['app_name'] = apk.get_app_name() or None
        analysis_result['target_sdk'] = apk.get_target_sdk_version() or None
        analysis_result['min_sdk'] = apk.get_min_sdk_version() or None
        analysis_result['max_sdk'] = apk.get_max_sdk_version() or None
        analysis_result['version_name'] = apk.get_androidversion_name() or None
        analysis_result['version_code'] = apk.get_androidversion_code() or None
        icon_data = apk.get_file(apk.get_app_icon())
        if icon_data:
            try:
                image = Image.open(io.BytesIO(icon_data))
                with io.BytesIO() as output:
                    image.save(output, format="PNG")
                    icon_base64 = base64.b64encode(output.getvalue()).decode('utf-8')
                    analysis_result['app_icon_base64'] = icon_base64
            except Exception as e:
                analysis_result['app_icon_base64'] = None
        else:
            analysis_result['app_icon_base64'] = None
        manifest_xml = apk.get_android_manifest_xml()
        root = ET.fromstring(ET.tostring(manifest_xml))

        analysis_result['exported_activities'] = [activity.get('{http://schemas.android.com/apk/res/android}name')
                                                  for activity in root.findall(".//activity") if is_exported(activity)]
        analysis_result['exported_services'] = [service.get('{http://schemas.android.com/apk/res/android}name')
                                                for service in root.findall(".//service") if is_exported(service)]
        analysis_result['exported_receivers'] = [receiver.get('{http://schemas.android.com/apk/res/android}name')
                                                 for receiver in root.findall(".//receiver") if is_exported(receiver)]
        analysis_result['exported_providers'] = [provider.get('{http://schemas.android.com/apk/res/android}name')
                                                 for provider in root.findall(".//provider") if is_exported(provider)]

        analysis_result['activities_length'] = len(analysis_result['activities'])
        analysis_result['services_length'] = len(analysis_result['services'])
        analysis_result['receivers_length'] = len(analysis_result['receivers'])
        analysis_result['providers_length'] = len(analysis_result['providers'])
        analysis_result['exported_activities_length'] = len(analysis_result['exported_activities'])
        analysis_result['exported_services_length'] = len(analysis_result['exported_services'])
        analysis_result['exported_receivers_length'] = len(analysis_result['exported_receivers'])
        analysis_result['exported_providers_length'] = len(analysis_result['exported_providers'])
        key_site = analyze_urls_in_apk(file_path)

        analysis_result['url_analysis_stats'] = key_site


        cert_info = get_v2_cert_info(apk)

        # 获取 v1 证书信息并合并
        #v1_cert_info = get_v1_cert_info(apk)
        #cert_info.extend(v1_cert_info)

        analysis_result['certificates'] = cert_info

        # 获取签名信息
        analysis_result['signatures'] = {
            'v1_signature': apk.is_signed_v1(),
            'v2_signature': apk.is_signed_v2(),
            'v3_signature': apk.is_signed_v3(),
            'is_signed': apk.is_signed(),
        }

        # 将值为0的字段显示为0而不是None
        for key in ['exported_activities_length', 'exported_services_length', 'exported_receivers_length', 'exported_providers_length']:
            if analysis_result[key] == 0:
                analysis_result[key] = 0

        # 过滤掉空值的键，但保留导出组件数量
        analysis_result = {key: value for key, value in analysis_result.items() if value is not None or key.endswith('_length')}
        api_calls = extract_api_calls(file_path)
        categorized_api_calls = categorize_apis(api_calls)
        analysis_result['android_apis'] = {k: v for k, v in categorized_api_calls.items() if
                                           v}

    except Exception as e:
        analysis_result['error'] = str(e)
    return analysis_result
def extract_urls_and_domains(apk_path):
    try:
        apk = APK(apk_path)
        urls = set()
        url_pattern = re.compile(r'https?://[^\s\'"<>]+')
        for dex in apk.get_all_dex():
            matches = re.findall(b'(https?://[^\s\'"<>]+)', dex)
            for match in matches:
                # 将多个URL分割开，并清除空格
                cleaned_urls = [url.strip() for url in re.split(r'(?=https?://)', match.decode('utf-8')) if url.strip()]
                urls.update(cleaned_urls)
        urls_and_domains = []
        for url in urls:
            domain = urlparse(url).netloc
            urls_and_domains.append({'url': url, 'domain': domain})
        urls_and_domains = [dict(t) for t in {tuple(d.items()) for d in urls_and_domains}]
        return urls_and_domains
    except Exception as e:
        print(f"Error processing {apk_path}: {e}")
        return []
def get_v2_cert_info(apk):
    certificates = apk.get_certificates_der_v2()
    cert_info = []
    if certificates:
        for cert in certificates:
            x509 = load_certificate(FILETYPE_ASN1, cert)
            pubkey = x509.get_pubkey()
            pubkey_algorithm_name = get_public_key_algorithm_name(pubkey.type())
            cert_info.append({
                'subject': format_cert_info(x509.get_subject().get_components()),
                'issuer': format_cert_info(x509.get_issuer().get_components()),
                'serial_number': hex(x509.get_serial_number()),
                'not_before': format_datetime(x509.get_notBefore().decode()),
                'not_after': format_datetime(x509.get_notAfter().decode()),
                'signature_algorithm': x509.get_signature_algorithm().decode(),
                'public_key_algorithm': pubkey_algorithm_name,

            })
    return cert_info
def get_v1_cert_info(apk):
    cert_info = []

    for cert in apk.get_certificates():
        cert_info.append({
            'hash_algorithm': cert.signature_hash_algorithm.name,
            'md5': cert.fingerprint(hashes.MD5()).hex(),
            'sha1': cert.fingerprint(hashes.SHA1()).hex(),
            'sha256': cert.fingerprint(hashes.SHA256()).hex(),
            'sha512': cert.fingerprint(hashes.SHA512()).hex(),
        })
    return cert_info

def format_datetime(dt_str):
    dt = datetime.strptime(dt_str, "%Y%m%d%H%M%SZ")
    return dt.strftime("%Y-%m-%d %H:%M:%S+00:00")

def format_cert_info(components):
    formatted_info = []
    for component in components:
        formatted_info.append(f"{component[0].decode()}={component[1].decode()}")
    return ",".join(formatted_info)

def calculate_hash(file_path, hash_type):
    hash_func = hashlib.new(hash_type)
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b''):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def is_exported(element):
    return element.get('{http://schemas.android.com/apk/res/android}exported') == 'true' or (
        element.get('{http://schemas.android.com/apk/res/android}exported') is None and element.get('{http://schemas.android.com/apk/res/android}permission') is None)

def format_size_to_mb(size_bytes):
    mb_size = size_bytes / (1024 * 1024)
    return f"{mb_size:.2f} MB"


# 添加一个映射函数来获取公钥类型的名称
def get_public_key_algorithm_name(pubkey_type):
    if pubkey_type == TYPE_RSA:
        return "RSA"
    elif pubkey_type == TYPE_DSA:
        return "DSA"
    else:
        return "Unknown"

def extract_api_calls(file_path):
    api_calls = set()
    try:
        a, d, dx = AnalyzeAPK(file_path)
        for method in dx.get_methods():
            for _, call, _ in method.get_xref_to():
                if call.class_name.startswith('Landroid/') or call.class_name.startswith('Ljava/'):
                    api_calls.add(f'{call.class_name}->{call.name}')
    except Exception as e:
        api_calls.add(f'Error extracting API calls: {str(e)}')
    return sorted(api_calls)

def categorize_apis(api_list):
    categories = {
        'Get System Service': [],
        'Inter Process Communication': [],
        'Java Reflection': [],
        'Local File I/O Operations': [],
        'Sending Broadcast': [],
        'Starting Activity': [],
        'WebView JavaScript Interface': [],
        'Content Providers': [],
        'Networking': [],
        'Location Services': [],
        'Bluetooth': [],
        'Sensors': [],
        'Media': [],
        'Database': [],
        'User Interface': [],
        'Permissions': [],
        'Security': [],
        'Analytics': [],
        'Push Notifications': [],
        'Graphics': [],
        'Animation': [],
        'Accessibility': [],
        'Input': [],
        'Telephony': [],
        'Storage': [],
        'Hardware': [],
        'Lifecycle': [],
        'Navigation': [],
        'Work': [],
        'Notifications': [],
        'Multithreading': [],
        'Clipboard': [],
        'Print': [],
        'Text': [],
        'Web': [],
        'App Widgets': [],
        'Fragments': [],
        'Resources': [],
        'Style and Themes': [],
        'Data Binding': [],
        'Kotlin Coroutines': [],
        'Jetpack Compose': []
    }

    for api in api_list:
        categorized = False
        for prefix, category in API_CATEGORIES.items():
            if api.startswith(prefix):
                if category not in categories:
                    categories[category] = []
                categories[category].append(api)
                categorized = True
                break

        if not categorized:
            if 'getSystemService' in api:
                categories['Get System Service'].append(api)
            elif 'startActivity' in api:
                categories['Starting Activity'].append(api)
            elif 'sendBroadcast' in api:
                categories['Sending Broadcast'].append(api)
            elif 'Class;' in api or 'Method;' in api or 'Field;' in api:
                categories['Java Reflection'].append(api)
            elif 'File;' in api or 'InputStream;' in api or 'OutputStream;' in api:
                categories['Local File I/O Operations'].append(api)
            elif 'WebView;' in api:
                categories['WebView JavaScript Interface'].append(api)
            elif 'ContentResolver;' in api or 'ContentProvider;' in api:
                categories['Content Providers'].append(api)
            elif 'HttpURLConnection;' in api or 'OkHttp;' in api:
                categories['Networking'].append(api)
            elif 'LocationManager;' in api or 'FusedLocationProviderClient;' in api:
                categories['Location Services'].append(api)
            elif 'BluetoothAdapter;' in api or 'BluetoothDevice;' in api:
                categories['Bluetooth'].append(api)
            elif 'SensorManager;' in api or 'Sensor;' in api:
                categories['Sensors'].append(api)
            elif 'MediaPlayer;' in api or 'MediaRecorder;' in api:
                categories['Media'].append(api)
            elif 'SQLiteDatabase;' in api or 'Room;' in api:
                categories['Database'].append(api)
            elif 'View;' in api or 'Button;' in api or 'RecyclerView;' in api:
                categories['User Interface'].append(api)
            elif 'checkSelfPermission' in api or 'requestPermissions' in api:
                categories['Permissions'].append(api)
            elif 'KeyStore;' in api or 'Cipher;' in api or 'Signature;' in api:
                categories['Security'].append(api)
            elif 'FirebaseAnalytics;' in api or 'GoogleAnalytics;' in api:
                categories['Analytics'].append(api)
            elif 'FirebaseMessaging;' in api or 'GCM;' in api:
                categories['Push Notifications'].append(api)
            elif 'Canvas;' in api or 'Bitmap;' in api or 'SurfaceView;' in api:
                categories['Graphics'].append(api)
            elif 'Animator;' in api or 'Animation;' in api or 'ObjectAnimator;' in api:
                categories['Animation'].append(api)
            elif 'AccessibilityManager;' in api or 'AccessibilityEvent;' in api:
                categories['Accessibility'].append(api)
            elif 'KeyEvent;' in api or 'MotionEvent;' in api or 'InputMethodManager;' in api:
                categories['Input'].append(api)
            elif 'TelephonyManager;' in api or 'SmsManager;' in api or 'SubscriptionManager;' in api:
                categories['Telephony'].append(api)
            elif 'SharedPreferences;' in api or 'StorageManager;' in api or 'DocumentFile;' in api:
                categories['Storage'].append(api)
            elif 'Camera;' in api or 'Sensor;' in api or 'Vibrator;' in api:
                categories['Hardware'].append(api)
            elif 'Activity;' in api or 'Fragment;' in api or 'LifecycleObserver;' in api:
                categories['Lifecycle'].append(api)
            elif 'NavController;' in api or 'NavHostFragment;' in api or 'DeepLink;' in api:
                categories['Navigation'].append(api)
            elif 'WorkManager;' in api or 'JobScheduler;' in api or 'AlarmManager;' in api:
                categories['Work'].append(api)
            elif 'NotificationManager;' in api or 'NotificationChannel;' in api or 'NotificationCompat;' in api:
                categories['Notifications'].append(api)
            elif 'AsyncTask;' in api or 'ThreadPoolExecutor;' in api or 'Handler;' in api:
                categories['Multithreading'].append(api)
            elif 'ClipboardManager;' in api or 'ClipData;' in api:
                categories['Clipboard'].append(api)
            elif 'PrintManager;' in api or 'PrintAttributes;' in api or 'PrintDocumentAdapter;' in api:
                categories['Print'].append(api)
            elif 'Editable;' in api or 'Spannable;' in api or 'TextWatcher;' in api:
                categories['Text'].append(api)
            elif 'WebView;' in api or 'CookieManager;' in api or 'WebResourceRequest;' in api:
                categories['Web'].append(api)
            elif 'AppWidgetManager;' in api or 'AppWidgetProvider;' in api or 'RemoteViews;' in api:
                categories['App Widgets'].append(api)
            elif 'FragmentManager;' in api or 'FragmentTransaction;' in api or 'FragmentActivity;' in api:
                categories['Fragments'].append(api)
            elif 'Resources;' in api or 'AssetManager;' in api or 'TypedArray;' in api:
                categories['Resources'].append(api)
            elif 'Style;' in api or 'Theme;' in api or 'TypedValue;' in api:
                categories['Style and Themes'].append(api)
            elif 'DataBindingUtil;' in api or 'ViewDataBinding;' in api or 'ObservableField;' in api:
                categories['Data Binding'].append(api)
            elif 'CoroutineScope;' in api or 'launch;' in api or 'async;' in api:
                categories['Kotlin Coroutines'].append(api)
            elif 'Composable;' in api or 'remember;' in api or 'Modifier;' in api:
                categories['Jetpack Compose'].append(api)
            else:
                continue

    return categories

def result_page(request):
    analysis_result = request.session.get('analysis_result', {})
    return render(request, 'static_ana.html', {'analysis_result': analysis_result})

# 生成静态分析的Word文件
def generate_file_static(analysis_result, file_name):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'SimSun')
    document.add_heading('静态与动态分析报告', 1)

    # 添加时间戳
    timestamp = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
    document.add_paragraph(f"生成时间: {timestamp}")

    # 添加分析结果的各个部分
    document.add_heading('文件信息', level=2)
    document.add_paragraph(f"文件名: {analysis_result.get('file_name', '无')}")
    document.add_paragraph(f"大小: {analysis_result.get('size', '无')}")
    document.add_paragraph(f"MD5: {analysis_result.get('md5', '无')}")
    document.add_paragraph(f"SHA1: {analysis_result.get('sha1', '无')}")
    document.add_paragraph(f"SHA256: {analysis_result.get('sha256', '无')}")

    document.add_heading('应用信息', level=2)
    document.add_paragraph(f"包名: {analysis_result.get('package', '无')}")
    document.add_paragraph(f"主活动: {analysis_result.get('main_activity', '')}")
    document.add_paragraph(f"应用名称: {analysis_result.get('app_name', '')}")
    document.add_paragraph(f"版本名称: {analysis_result.get('version_name', '')}")
    document.add_paragraph(f"版本代码: {analysis_result.get('version_code', '')}")

    # 权限
    document.add_heading('权限', level=2)
    permissions = analysis_result.get('permissions', [])
    if permissions:
        for perm in permissions:
            document.add_paragraph(f"{perm['Permission']}: {perm['Description']}")
    else:
        document.add_paragraph("无")

    # 活动
    document.add_heading('活动', level=2)
    activities = analysis_result.get('activities', [])
    if activities:
        for activity in activities:
            document.add_paragraph(activity)
    else:
        document.add_paragraph("无")

    # 服务
    document.add_heading('服务', level=2)
    services = analysis_result.get('services', [])
    if services:
        for service in services:
            document.add_paragraph(service)
    else:
        document.add_paragraph("无")

    # 接收器
    document.add_heading('接收器', level=2)
    receivers = analysis_result.get('receivers', [])
    if receivers:
        for receiver in receivers:
            document.add_paragraph(receiver)
    else:
        document.add_paragraph("无")

    # 提供者
    document.add_heading('提供者', level=2)
    providers = analysis_result.get('providers', [])
    if providers:
        for provider in providers:
            document.add_paragraph(provider)
    else:
        document.add_paragraph("无")

    # 证书
    document.add_heading('证书', level=2)
    certificates = analysis_result.get('certificates', [])
    if certificates:
        for cert in certificates:
            document.add_paragraph(f"主题: {cert.get('subject', '')}")
            document.add_paragraph(f"颁发者: {cert.get('issuer', '')}")
            document.add_paragraph(f"序列号: {cert.get('serial_number', '')}")
            document.add_paragraph(f"有效期起: {cert.get('not_before', '')}")
            document.add_paragraph(f"有效期至: {cert.get('not_after', '')}")
            document.add_paragraph(f"签名算法: {cert.get('signature_algorithm', '')}")
            document.add_paragraph(f"公钥算法: {cert.get('public_key_algorithm', '')}")
    else:
        document.add_paragraph("无")
    document.add_heading('API调用', level=2)
    if analysis_result.get('android_apis'):
        for category, apis in analysis_result['android_apis'].items():
            document.add_heading(category, level=3)
            document.add_paragraph(', '.join(apis) if apis else '无')
    else:
        document.add_paragraph("无")

    document.add_heading('关键站点', level=2)
    key_sites = analysis_result.get('url_analysis_stats', [])
    if key_sites:
        for site in key_sites:
            document.add_paragraph(f"URL: {site.get('url', '无')}")
            document.add_paragraph(f"域名: {site.get('domain', '无')}")
            document.add_paragraph(f"危险等级: {site.get('danger_level', '无')}")
            document.add_paragraph(f"归一化危险等级: {site.get('normalized_danger_level', '无')}")
    else:
        document.add_paragraph("无")

    # 生成文件路径
    report_dir = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0])
    os.makedirs(report_dir, exist_ok=True)
    word_file_path = os.path.join(report_dir, f'{file_name}_static_and_dynamic_analysis.docx')
    document.save(word_file_path)
    pdf_file_path = word_file_path.replace('.docx', '.pdf')
    convert(word_file_path, pdf_file_path)
    return word_file_path,pdf_file_path


def generate_file_model(prediction_result, file_name):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'SimSun')
    document.add_heading('模型预测报告', 1)
    timestamp = prediction_result.get('scan_time', '')
    document.add_paragraph(f"预测时间: {timestamp}")
    document.add_heading('预测结果', level=2)
    document.add_paragraph(f"文件名: {prediction_result.get('file_name', '')}")
    prediction = prediction_result.get('prediction', '')
    confidence = prediction_result.get('confidence', '未知')
    if prediction == 0:
        prediction_text = "WHITE"
    elif prediction == 1:
        prediction_text = "SCAM"
    else:
        prediction_text = "Unknown"
    document.add_paragraph(f"集成学习预测: {prediction_text}")
    document.add_paragraph(f"集成学习置信度: {confidence:.5f}")
    # 添加GCN预测结果
    gcn_prediction = prediction_result.get('gcn_prediction', '未知')
    gcn_confidence = prediction_result.get('gcn_confidence', '未知')

    if gcn_prediction == 0:
        gcn_prediction_text = "WHITE"
    elif gcn_prediction == 1:
        gcn_prediction_text = "SCAM"
    else:
        gcn_prediction_text = "Unknown"
    document.add_paragraph(f"GCN预测: {gcn_prediction_text}")
    document.add_paragraph(f"GCN置信度: {gcn_confidence:.5f}")
    # 将HTML转换为图像
    html_file_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0],
                                  'lime_explanation.html')
    image_file_path = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0],
                                   'lime_explanation.jpg')
    convert_html_to_image(html_file_path, image_file_path)
    if os.path.exists(image_file_path):
        document.add_heading('LIME Explanation', level=2)
        document.add_picture(image_file_path, width=Inches(6))
    # 生成文件路径
    report_dir = os.path.join(settings.MEDIA_ROOT, 'apk_files', os.path.splitext(file_name)[0])
    os.makedirs(report_dir, exist_ok=True)
    word_file_path = os.path.join(report_dir, f'{file_name}_model_prediction.docx')
    document.save(word_file_path)
    pdf_file_path = word_file_path.replace('.docx', '.pdf')
    convert(word_file_path, pdf_file_path)
    return word_file_path,pdf_file_path
def download_report_zip(request, apk_file_name):
    apk = get_object_or_404(ApkFile, file_name=apk_file_name)
    report_files=[]
    if(apk.static_analyzed):
        report_files.append(apk.word_report_path_static)
        report_files.append(apk.pdf_report_path_static)
    if(apk.prediction_result is not None):
        report_files.append(apk.pdf_report_path_model)
        report_files.append(apk.word_report_path_model)
    zip_subdir = "报告"
    zip_filename = f"{zip_subdir}.zip"
    zip_path = os.path.join(settings.MEDIA_ROOT, zip_filename)

    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file_path in report_files:
            if file_path and os.path.exists(file_path):
                file_dir, file_name = os.path.split(file_path)
                zipf.write(file_path, os.path.join(zip_subdir, file_name))

    # 读取zip文件并返回响应
    response = HttpResponse(open(zip_path, 'rb'), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename={zip_filename}'
    return response
def static_export_report(request):
    if request.method == 'POST':
        format = request.POST.get('format')
        apk_file_name = request.POST.get('apk_file_name')
        apk_record = get_object_or_404(ApkFile, file_name=apk_file_name)

        if format == 'pdf':
            file_path = apk_record.pdf_report_path_static
        elif format == 'docx':
            file_path = apk_record.word_report_path_static
        else:
            return JsonResponse({'status': 'error', 'message': 'Invalid format'}, status=400)

        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'status': 'error', 'message': 'File not found'}, status=404)

        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=os.path.basename(file_path))
        return response

    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)
def convert_html_to_image(html_file_path, image_file_path):
    try:
        path_wkimg = os.path.join(settings.WKHTML_ROOT, 'wkhtmltopdf', 'bin', 'wkhtmltoimage.exe')
        path_wkimg = os.path.normpath(path_wkimg)  # 标准化路径
        cfg = imgkit.config(wkhtmltoimage=path_wkimg)
        options = {
            'javascript-delay': '2000',
            'height': '800',
            'width': '2000'
        }
        html_file_path = os.path.normpath(html_file_path)
        image_file_path = os.path.normpath(image_file_path)

        imgkit.from_file(html_file_path, image_file_path, config=cfg, options=options)
    except Exception as e:
        print(f"Error converting HTML to image: {e}")
        raise e  # 重新抛出异常以便外部捕获
